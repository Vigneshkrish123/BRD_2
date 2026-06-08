"""
formatter.py — BRD JSON → styled .docx (Polycab house style).

Renders the v2 use-case-driven schema produced by generator.py into a Word
document that mirrors the reference BRDs: title page, numbered sections,
navy-header tables, In/Out-of-Scope tables, AS IS / TO BE flows, and full
use-case blocks (Description, Role/Action/Benefit, Pre/Post-Condition,
Main Flow table, Business Rules, Exceptional Flow, per-UC Out of Scope).

Pure python-docx. No Node, no subprocess.

Public entry point:
    build_brd(brd: dict, output_path: str) -> str
"""

from __future__ import annotations

import io
import tempfile
import os
from datetime import date

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Inches

# ── Palette (matches reference: Polycab navy / blue) ──────────────────────────
NAVY = RGBColor(0x1F, 0x38, 0x64)        # heading 2 / table header fill
BLUE = RGBColor(0x2E, 0x75, 0xB6)        # heading 1
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
HEADER_FILL = "1F3864"                    # table header bg (hex, no #)
ROW_ALT_FILL = "EAF1F8"                   # zebra stripe
FONT = "Arial"


# ── Low-level helpers ─────────────────────────────────────────────────────────

def _set_cell_bg(cell, hex_fill: str) -> None:
    """Set table cell shading — python-docx has no native API for this."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_fill)
    tc_pr.append(shd)


def _set_cell_margins(cell, top=60, bottom=60, left=110, right=110) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    m = OxmlElement("w:tcMar")
    for edge, val in (("top", top), ("bottom", bottom), ("start", left), ("end", right)):
        node = OxmlElement(f"w:{edge}")
        node.set(qn("w:w"), str(val))
        node.set(qn("w:type"), "dxa")
        m.append(node)
    tc_pr.append(m)


def _style_cell_text(cell, *, bold=False, color: RGBColor | None = None, size=10):
    for para in cell.paragraphs:
        para.paragraph_format.space_before = Pt(1)
        para.paragraph_format.space_after = Pt(1)
        for run in para.runs:
            run.font.name = FONT
            run.font.size = Pt(size)
            run.font.bold = bold
            if color is not None:
                run.font.color.rgb = color
        if not para.runs:  # empty cell still needs a styled (empty) run target
            run = para.add_run("")
            run.font.name = FONT
            run.font.size = Pt(size)


def _add_table(doc, headers: list[str], rows: list[list[str]], widths: list[float] | None = None):
    """Styled table: navy header, white header text, zebra body rows."""
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = False

    hdr = table.rows[0].cells
    for i, text in enumerate(headers):
        hdr[i].text = str(text)
        _set_cell_bg(hdr[i], HEADER_FILL)
        _set_cell_margins(hdr[i])
        _style_cell_text(hdr[i], bold=True, color=WHITE)

    for r, row in enumerate(rows):
        cells = table.add_row().cells
        for i in range(len(headers)):
            cells[i].text = str(row[i]) if i < len(row) else ""
            if r % 2 == 1:
                _set_cell_bg(cells[i], ROW_ALT_FILL)
            _set_cell_margins(cells[i])
            _style_cell_text(cells[i])

    if widths:
        for i, w in enumerate(widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)
    return table


def _heading(doc, text: str, level: int):
    p = doc.add_heading(level=level)
    run = p.add_run(text)
    run.font.name = FONT
    run.font.bold = True
    if level == 1:
        run.font.size = Pt(15)
        run.font.color.rgb = BLUE
    else:
        run.font.size = Pt(13)
        run.font.color.rgb = NAVY
    return p


def _para(doc, text: str, *, bold=False, italic=False, size=10.5, space_after=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.font.name = FONT
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    return p


def _label_value(doc, label: str, value: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    b = p.add_run(f"{label}: ")
    b.font.name = FONT
    b.font.size = Pt(10.5)
    b.font.bold = True
    v = p.add_run(value)
    v.font.name = FONT
    v.font.size = Pt(10.5)
    return p


def _bullet(doc, text: str, *, lead: str | None = None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    if lead:
        b = p.add_run(f"{lead}: ")
        b.font.name = FONT
        b.font.size = Pt(10.5)
        b.font.bold = True
    run = p.add_run(text)
    run.font.name = FONT
    run.font.size = Pt(10.5)
    return p


def _numbered(doc, text: str):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.font.name = FONT
    run.font.size = Pt(10.5)
    return p


def _join(values) -> str:
    """Render a list of pre/post-conditions as '1) a  2) b'."""
    if not values:
        return "—"
    if len(values) == 1:
        return str(values[0])
    return "  ".join(f"{i+1}) {v}" for i, v in enumerate(values))


# ── Section renderers ─────────────────────────────────────────────────────────

def _title_page(doc, info: dict):
    for _ in range(3):
        doc.add_paragraph()
    for text, size, bold in (
        ("Polycab India Limited.", 16, True),
        ("Business Requirements Document", 20, True),
    ):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        r.font.name = FONT
        r.font.size = Pt(size)
        r.font.bold = bold
        if text.startswith("Business"):
            r.font.color.rgb = NAVY

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"Project Name: {info.get('project_name', 'TBD')}")
    r.font.name = FONT
    r.font.size = Pt(16)
    r.font.bold = True
    r.font.color.rgb = BLUE

    for _ in range(4):
        doc.add_paragraph()

    meta = [
        ["Version", info.get("version", "1.0")],
        ["Status", info.get("status", "Draft")],
        ["Date", date.today().strftime("%d %B %Y")],
        ["Prepared By", info.get("prepared_by", "BRD Agent (AI-assisted)")],
    ]
    t = _add_table(doc, ["Field", "Details"], meta, widths=[2.0, 4.5])
    t.alignment = WD_TABLE_ALIGNMENT.CENTER

    doc.add_page_break()


def _introduction(doc, brd: dict):
    _heading(doc, "1. Introduction", 1)
    if brd.get("introduction"):
        _para(doc, brd["introduction"], size=10.5, space_after=8)

    _heading(doc, "1.1 Business Objectives", 2)
    for bo in brd.get("business_objectives", []):
        if isinstance(bo, dict):
            title = bo.get("title") or bo.get("id", "")
            _para(doc, title, bold=True, space_after=1)
            _para(doc, bo.get("description", ""), space_after=6)
        else:
            _bullet(doc, str(bo))

    if brd.get("stakeholders"):
        _heading(doc, "1.2 Stakeholders", 2)
        rows = [[s.get("role", ""), s.get("responsibility", "")] for s in brd["stakeholders"]]
        _add_table(doc, ["Role", "Responsibility"], rows, widths=[2.5, 4.0])


def _scope(doc, brd: dict):
    scope = brd.get("scope", {})
    _heading(doc, "2. Scope", 1)

    _heading(doc, "2.1 In-Scope", 2)
    in_scope = scope.get("in_scope", [])
    if in_scope and isinstance(in_scope[0], dict):
        rows = [[i.get("module", ""), i.get("feature", ""),
                 i.get("description", ""), i.get("key_outcomes", "")] for i in in_scope]
        _add_table(doc, ["Module", "Feature", "Description", "Key Outcomes"],
                   rows, widths=[1.3, 1.6, 2.4, 1.6])
    else:
        for item in in_scope:
            _bullet(doc, str(item))

    _heading(doc, "2.2 Out of Scope", 2)
    out_scope = scope.get("out_of_scope", [])
    if out_scope and isinstance(out_scope[0], dict):
        rows = [[o.get("item", ""), o.get("description", "")] for o in out_scope]
        _add_table(doc, ["Item", "Description"], rows, widths=[2.5, 4.0])
    else:
        for item in out_scope:
            _bullet(doc, str(item))


def _assumptions(doc, brd: dict):
    rows_src = brd.get("assumptions", [])
    if not rows_src:
        return
    _heading(doc, "3. Assumptions", 1)
    if isinstance(rows_src[0], dict):
        rows = [[str(i + 1), a.get("assumption", ""), a.get("impact_if_changed", "") or "—"]
                for i, a in enumerate(rows_src)]
        _add_table(doc, ["Sr. No.", "Assumption", "Impact If Changed"],
                   rows, widths=[0.9, 3.3, 2.3])
    else:
        for a in rows_src:
            _bullet(doc, str(a))


def _as_is(doc, brd: dict):
    flow = brd.get("as_is_business_flow", [])
    if not flow:
        return
    _heading(doc, "4. AS IS Business Flow", 1)
    for step in flow:
        _bullet(doc, str(step))


def _sub_bullet(doc, text: str):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.85)
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run(str(text))
    run.font.name = FONT
    run.font.size = Pt(10.5)
    return p


def _example(doc, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.85)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(f"Example: {text}")
    run.font.name = FONT
    run.font.size = Pt(10)
    run.font.italic = True
    run.font.color.rgb = NAVY
    return p


def _to_be(doc, brd: dict):
    flow = brd.get("to_be_business_process_flow", [])
    if not flow:
        return
    _heading(doc, "5. TO BE Business Process Flow", 1)
    for item in flow:
        if isinstance(item, dict):
            _numbered(doc, item.get("step", ""))
            for sub in item.get("sub_steps", []):
                _sub_bullet(doc, sub)
            if item.get("example"):
                _example(doc, item["example"])
        else:
            _numbered(doc, str(item))


def _use_cases(doc, brd: dict):
    ucs = brd.get("use_cases", [])
    if not ucs:
        return
    _heading(doc, "6. Use Cases", 1)
    for idx, uc in enumerate(ucs, start=1):
        uc_id = uc.get("id", f"UC_{idx:02d}")
        _heading(doc, f"6.{idx} {uc_id}: {uc.get('title', '')}", 2)

        _para(doc, "Description:", bold=True, space_after=1)
        _para(doc, uc.get("description", ""), space_after=6)

        if uc.get("role"):
            _bullet(doc, uc["role"], lead="Role")
        if uc.get("action"):
            _bullet(doc, uc["action"], lead="Action")
        if uc.get("benefit"):
            _bullet(doc, uc["benefit"], lead="Benefit")

        if uc.get("end_user"):
            _label_value(doc, "End User", uc["end_user"])
        _label_value(doc, "Pre-Condition", _join(uc.get("pre_conditions", [])))
        _label_value(doc, "Post-Condition", _join(uc.get("post_conditions", [])))

        # Main Flow — collapse User Action column when every step is system-driven.
        main_flow = uc.get("main_flow", [])
        if main_flow:
            _para(doc, "Main Flow:", bold=True, space_after=2)
            has_user = any((s.get("user_action") or "").strip() for s in main_flow)
            if has_user:
                rows = [[s.get("step", str(i + 1)), s.get("user_action", ""), s.get("system_action", "")]
                        for i, s in enumerate(main_flow)]
                _add_table(doc, ["Step", "User Action", "System Action"],
                           rows, widths=[0.7, 2.75, 3.05])
            else:
                rows = [[s.get("step", str(i + 1)), s.get("system_action", "")]
                        for i, s in enumerate(main_flow)]
                _add_table(doc, ["Step", "System Action"], rows, widths=[0.9, 5.6])

        if uc.get("business_rules"):
            _para(doc, "Business Rules:", bold=True, space_after=2)
            rows = [[str(i + 1), r] for i, r in enumerate(uc["business_rules"])]
            _add_table(doc, ["Sr. No.", "Business Rule"], rows, widths=[0.9, 5.6])

        ex = uc.get("exceptional_flow", [])
        if ex:
            _para(doc, "Exceptional Flow:", bold=True, space_after=2)
            rows = [[str(i + 1), e.get("exception", ""), e.get("error_message", "") or "—"]
                    for i, e in enumerate(ex)]
            _add_table(doc, ["Sr. No.", "Exception", "Error Message / Behaviour"],
                       rows, widths=[0.9, 2.8, 2.8])

        if uc.get("out_of_scope"):
            _para(doc, "Out of Scope:", bold=True, space_after=2)
            rows = [[str(i + 1), o] for i, o in enumerate(uc["out_of_scope"])]
            _add_table(doc, ["Sr. No.", "Description"], rows, widths=[0.9, 5.6])

        doc.add_paragraph()


def _nfr(doc, brd: dict, n: int):
    nfrs = brd.get("non_functional_requirements", [])
    if not nfrs:
        return n
    _heading(doc, f"{n}. Non-Functional Requirements", 1)
    if isinstance(nfrs[0], dict):
        rows = [[x.get("id", f"NFR-{i+1:03d}"), x.get("category", ""),
                 x.get("description", ""), x.get("priority", "")] for i, x in enumerate(nfrs)]
        _add_table(doc, ["ID", "Category", "Requirement", "Priority"],
                   rows, widths=[1.0, 1.4, 3.2, 0.9])
    else:
        for x in nfrs:
            _bullet(doc, str(x))
    return n + 1


def _risks(doc, brd: dict, n: int):
    risks = brd.get("risks", [])
    if not risks:
        return n
    _heading(doc, f"{n}. Risks", 1)
    rows = [[r.get("id", f"R-{i+1:03d}"), r.get("description", ""),
             r.get("impact", ""), r.get("mitigation", "")] for i, r in enumerate(risks)]
    _add_table(doc, ["ID", "Risk", "Impact", "Mitigation"], rows, widths=[0.9, 2.6, 0.9, 2.1])
    return n + 1


def _open_questions(doc, brd: dict, n: int):
    oqs = brd.get("open_questions", [])
    if not oqs:
        return n
    _heading(doc, f"{n}. Open Questions", 1)
    rows = [[q.get("id", f"OQ-{i+1:03d}"), q.get("question", ""),
             q.get("owner", ""), q.get("target_date", "")] for i, q in enumerate(oqs)]
    _add_table(doc, ["ID", "Question", "Owner", "Target Date"], rows, widths=[0.9, 3.4, 1.3, 0.9])
    return n + 1


# ── Document-level setup ──────────────────────────────────────────────────────

def _setup_base_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal.font.size = Pt(10.5)


def _add_footer_page_numbers(doc):
    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    fld_begin = OxmlElement("w:fldChar"); fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve"); instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar"); fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin); run._r.append(instr); run._r.append(fld_end)
    run.font.name = FONT
    run.font.size = Pt(9)


# ── Public entry point ────────────────────────────────────────────────────────

def build_brd(brd: dict, output_path: str) -> str:
    doc = Document()
    _setup_base_styles(doc)
    _add_footer_page_numbers(doc)

    info = brd.get("document_info", {})
    _title_page(doc, info)

    _introduction(doc, brd)
    _scope(doc, brd)
    _assumptions(doc, brd)
    _as_is(doc, brd)
    _to_be(doc, brd)
    _use_cases(doc, brd)

    n = 7
    n = _nfr(doc, brd, n)
    n = _risks(doc, brd, n)
    n = _open_questions(doc, brd, n)

    doc.save(output_path)
    return output_path


def format_docx(brd: dict) -> bytes:
    """Convenience wrapper: build BRD and return raw .docx bytes (for Streamlit download)."""
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        tmp_path = tmp.name
    try:
        build_brd(brd, tmp_path)
        with open(tmp_path, "rb") as f:
            return f.read()
    finally:
        os.unlink(tmp_path)
